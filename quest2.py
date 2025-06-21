import streamlit as st
import os
import io
import json
import pandas as pd
import docx
import PyPDF2
import requests
import traceback
import re
from datetime import datetime
import hashlib
from typing import Dict, List, Any, Tuple, Optional
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
import tiktoken
import plotly.express as px
import plotly.graph_objects as go

# LangChain imports
from langchain.llms import OpenAI
from langchain.chat_models import ChatOpenAI, ChatAnthropic
from langchain.llms.base import LLM
from langchain.schema import BaseMessage, HumanMessage, SystemMessage
from langchain.callbacks.manager import CallbackManagerForLLMRun
from langchain.schema.output import LLMResult
from langchain.prompts import PromptTemplate, ChatPromptTemplate
from langchain.output_parsers import PydanticOutputParser, OutputFixingParser
from langchain.schema import OutputParserException
from pydantic import BaseModel, Field
import yaml

# Download required NLTK data (run once)
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
except LookupError:
    with st.spinner("Downloading language models..."):
        nltk.download('punkt', quiet=True)
        nltk.download('stopwords', quiet=True)

# Enhanced Question Types
QUESTION_TYPES = [
    ('MCQ', 'Multiple Choice Question'),
    ('FIB', 'Fill in the Blanks'),
    ('SAQ', 'Short Answer Question'),
    ('LAQ', 'Long Answer Question'),
    ('TRUE_FALSE', 'True/False Question'),
    ('MATCHING', 'Matching'),
    ('ORDERING', 'Ordering/Sequencing'),
    ('CLASSIFICATION', 'Classification'),
    ('CASE_STUDY', 'Case Study/Scenario-Based'),
]

# Bloom's Taxonomy Levels with detailed descriptions
BLOOMS_TAXONOMY = {
    'remember': {
        'level': 1,
        'description': 'Recall facts, basic concepts, and answers',
        'verbs': ['define', 'list', 'identify', 'name', 'state', 'describe', 'match', 'select'],
        'question_starters': ['What is...?', 'Who was...?', 'When did...?', 'Where is...?', 'List the...', 'Define...']
    },
    'understand': {
        'level': 2,
        'description': 'Explain ideas and concepts',
        'verbs': ['explain', 'interpret', 'summarize', 'paraphrase', 'classify', 'compare', 'contrast'],
        'question_starters': ['Explain why...', 'What is the main idea of...?', 'Summarize...', 'Compare...', 'What does... mean?']
    },
    'apply': {
        'level': 3,
        'description': 'Use information in new situations',
        'verbs': ['apply', 'demonstrate', 'calculate', 'solve', 'show', 'complete', 'use', 'examine'],
        'question_starters': ['How would you use...?', 'What examples can you find...?', 'How would you solve...?', 'Apply the concept...']
    },
    'analyze': {
        'level': 4,
        'description': 'Draw connections among ideas',
        'verbs': ['analyze', 'categorize', 'examine', 'compare', 'contrast', 'separate', 'distinguish'],
        'question_starters': ['What are the parts of...?', 'How does... relate to...?', 'Why do you think...?', 'What evidence supports...?']
    },
    'evaluate': {
        'level': 5,
        'description': 'Justify a stand or decision',
        'verbs': ['evaluate', 'judge', 'critique', 'assess', 'defend', 'support', 'argue', 'recommend'],
        'question_starters': ['Which is better...?', 'What is your opinion of...?', 'How would you prioritize...?', 'What criteria would you use...?']
    },
    'create': {
        'level': 6,
        'description': 'Produce new or original work',
        'verbs': ['create', 'design', 'formulate', 'build', 'invent', 'develop', 'compose', 'generate'],
        'question_starters': ['How would you design...?', 'What would happen if...?', 'Can you create...?', 'How would you improve...?']
    }
}

# Pydantic models for structured output
class QuestionModel(BaseModel):
    id: str = Field(description="Unique identifier for the question")
    difficulty: str = Field(description="Difficulty level: easy, medium, or hard")
    difficulty_score: int = Field(description="Difficulty score from 1-10")
    type: str = Field(description="Question type code (MCQ, FIB, SAQ, etc.)")
    type_description: str = Field(description="Full question type description")
    cognitive_level: str = Field(description="Bloom's taxonomy level")
    cognitive_description: str = Field(description="Description of cognitive level")
    question: str = Field(description="The question text")
    options: Optional[List[str]] = Field(default=[], description="Multiple choice options if applicable")
    correct_answer: str = Field(description="The correct answer")
    explanation: str = Field(description="Detailed explanation")
    keywords: List[str] = Field(default=[], description="Key terms and concepts")
    estimated_time: int = Field(description="Estimated time in minutes")
    points: int = Field(description="Point value for the question")
    prerequisites: List[str] = Field(default=[], description="Required prerequisite knowledge")
    common_mistakes: List[str] = Field(default=[], description="Common student errors")
    follow_up: str = Field(default="", description="Follow-up question suggestion")
    # Additional fields for new question types
    matching_pairs: Optional[List[Dict[str, str]]] = Field(default=[], description="Matching pairs for matching questions")
    sequence_items: Optional[List[str]] = Field(default=[], description="Items to be ordered for sequencing questions")
    categories: Optional[List[str]] = Field(default=[], description="Categories for classification questions")
    scenario: Optional[str] = Field(default="", description="Scenario for case study questions")

class QuestionSetModel(BaseModel):
    questions: List[QuestionModel] = Field(description="List of generated questions")
    metadata: Dict[str, Any] = Field(description="Generation metadata")

# Custom LangChain LLM for DeepSeek
class DeepSeekLLM(LLM):
    api_key: str
    model: str = "deepseek-chat"
    base_url: str = "https://api.deepseek.com/v1"
    temperature: float = 0.7
    max_tokens: int = 3000
    
    @property
    def _llm_type(self) -> str:
        return "deepseek"
    
    def _call(
        self,
        prompt: str,
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs: Any,
    ) -> str:
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": self.model,
            "messages": [
                {"role": "system", "content": "You are an expert educational assessment creator specializing in Bloom's Taxonomy. Respond only with valid JSON."},
                {"role": "user", "content": prompt}
            ],
            "temperature": self.temperature,
            "max_tokens": self.max_tokens
        }
        
        response = requests.post(f"{self.base_url}/chat/completions", headers=headers, json=data)
        response.raise_for_status()
        
        response_data = response.json()
        return response_data['choices'][0]['message']['content']

# Custom LangChain LLM for other providers
class CustomAPILLM(LLM):
    api_key: str
    model: str
    base_url: str
    provider_name: str
    temperature: float = 0.7
    max_tokens: int = 3000
    
    @property
    def _llm_type(self) -> str:
        return self.provider_name.lower()
    
    def _call(
        self,
        prompt: str,
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs: Any,
    ) -> str:
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": self.model,
            "messages": [
                {"role": "system", "content": "You are an expert educational assessment creator specializing in Bloom's Taxonomy. Respond only with valid JSON."},
                {"role": "user", "content": prompt}
            ],
            "temperature": self.temperature,
            "max_tokens": self.max_tokens
        }
        
        response = requests.post(f"{self.base_url}/chat/completions", headers=headers, json=data)
        response.raise_for_status()
        
        response_data = response.json()
        return response_data['choices'][0]['message']['content']

class BloomsTaxonomyQuestionGenerator:
    """Enhanced question generator with Bloom's Taxonomy focus"""
    
    def __init__(self):
        self.model_provider = LangChainModelProvider()
        self.max_tokens = 3000
        self.question_history = []
        self.analytics_data = []
        
        # Question templates based on Bloom's Taxonomy and question types
        self.bloom_question_templates = {
            'remember': {
                'MCQ': "What is the definition of {concept}?",
                'FIB': "Fill in the blank: {concept} is defined as ______.",
                'SAQ': "Define {concept} in your own words.",
                'TRUE_FALSE': "{statement} is true or false?",
                'MATCHING': "Match the following {concepts} with their definitions.",
            },
            'understand': {
                'MCQ': "Which statement best explains {concept}?",
                'SAQ': "Explain the main idea behind {concept}.",
                'LAQ': "Describe how {concept} works and its significance.",
                'CLASSIFICATION': "Classify the following examples into appropriate categories of {concept}.",
            },
            'apply': {
                'MCQ': "In which scenario would you use {concept}?",
                'SAQ': "How would you apply {concept} to solve {problem}?",
                'LAQ': "Demonstrate the application of {concept} in a real-world situation.",
                'CASE_STUDY': "Given the following scenario, apply {concept} to analyze the situation.",
            },
            'analyze': {
                'MCQ': "What is the relationship between {concept1} and {concept2}?",
                'SAQ': "Analyze the components of {concept}.",
                'LAQ': "Break down {concept} into its constituent parts and explain their relationships.",
                'CLASSIFICATION': "Analyze and categorize the following elements based on {criteria}.",
            },
            'evaluate': {
                'MCQ': "Which approach to {concept} is most effective?",
                'SAQ': "Evaluate the strengths and weaknesses of {concept}.",
                'LAQ': "Critically assess {concept} and justify your evaluation.",
                'CASE_STUDY': "Evaluate the effectiveness of the solution in the given scenario.",
            },
            'create': {
                'SAQ': "Design a new approach to {problem} using {concept}.",
                'LAQ': "Create a comprehensive plan that incorporates {concept}.",
                'CASE_STUDY': "Develop a solution for the given scenario using creative thinking.",
                'ORDERING': "Create a logical sequence for implementing {process}.",
            }
        }
    
    def initialize_providers(self, provider_configs: Dict) -> Dict[str, str]:
        """Initialize AI model providers using LangChain"""
        errors = {}
        
        for provider, config in provider_configs.items():
            if config.get('api_key'):
                success = self.model_provider.initialize_model(provider, config)
                if not success:
                    errors[provider] = "Failed to initialize"
        
        return errors
    
    def extract_text_from_pdf(self, file) -> str:
        """Enhanced PDF extraction with better error handling"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            total_pages = len(pdf_reader.pages)
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for page_num in range(total_pages):
                try:
                    page_text = pdf_reader.pages[page_num].extract_text()
                    page_text = re.sub(r'\n+', '\n', page_text)
                    page_text = re.sub(r'\s+', ' ', page_text)
                    page_text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\"\'\/]', '', page_text)
                    text += page_text + "\n"
                    
                    progress = (page_num + 1) / total_pages
                    progress_bar.progress(progress)
                    status_text.text(f"Processing page {page_num + 1} of {total_pages}")
                    
                except Exception as e:
                    st.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                    continue
            
            progress_bar.empty()
            status_text.empty()
            return text.strip()
            
        except Exception as e:
            st.error(f"Error reading PDF: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, file) -> str:
        """Enhanced DOCX extraction with tables and formatting"""
        try:
            doc = docx.Document(file)
            text = ""
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text.strip() + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            return text.strip()
        except Exception as e:
            st.error(f"Error reading DOCX: {str(e)}")
            return ""
    
    def extract_content(self, uploaded_file) -> str:
        """Extract content from various file types"""
        file_type = uploaded_file.name.split('.')[-1].lower()
        
        try:
            if file_type == 'pdf':
                return self.extract_text_from_pdf(uploaded_file)
            elif file_type in ['docx', 'doc']:
                return self.extract_text_from_docx(uploaded_file)
            elif file_type in ['xlsx', 'xls']:
                df = pd.read_excel(uploaded_file)
                return df.to_string(index=False)
            elif file_type == 'txt':
                return uploaded_file.getvalue().decode('utf-8')
            elif file_type == 'csv':
                df = pd.read_csv(uploaded_file)
                return df.to_string(index=False)
            else:
                return f"Unsupported file type: {file_type}"
        except Exception as e:
            return f"Error extracting content: {str(e)}"
    
    def intelligent_content_chunking(self, content: str, max_chunk_tokens: int = 2000) -> List[str]:
        """Intelligently chunk content by sentences and paragraphs"""
        paragraphs = content.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
                
            test_chunk = current_chunk + "\n\n" + paragraph if current_chunk else paragraph
            
            if self.model_provider.count_tokens(test_chunk) <= max_chunk_tokens:
                current_chunk = test_chunk
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                
                if self.model_provider.count_tokens(paragraph) > max_chunk_tokens:
                    sentences = sent_tokenize(paragraph)
                    sentence_chunk = ""
                    
                    for sentence in sentences:
                        test_sentence_chunk = sentence_chunk + " " + sentence if sentence_chunk else sentence
                        
                        if self.model_provider.count_tokens(test_sentence_chunk) <= max_chunk_tokens:
                            sentence_chunk = test_sentence_chunk
                        else:
                            if sentence_chunk:
                                chunks.append(sentence_chunk)
                            sentence_chunk = sentence
                    
                    if sentence_chunk:
                        current_chunk = sentence_chunk
                else:
                    current_chunk = paragraph
        
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def extract_key_topics(self, content: str) -> List[str]:
        """Extract key topics using enhanced NLP techniques"""
        try:
            words = word_tokenize(content.lower())
            stop_words = set(stopwords.words('english'))
            
            filtered_words = [
                word for word in words 
                if word.isalpha() and word not in stop_words and len(word) > 3
                and not word.isdigit()
            ]
            
            word_freq = {}
            for word in filtered_words:
                word_freq[word] = word_freq.get(word, 0) + 1
            
            phrases = []
            words_list = content.split()
            for i in range(len(words_list) - 1):
                phrase = f"{words_list[i]} {words_list[i+1]}".lower()
                if len(phrase) > 6 and phrase.count(' ') == 1:
                    phrases.append(phrase)
            
            phrase_freq = {}
            for phrase in phrases:
                phrase_freq[phrase] = phrase_freq.get(phrase, 0) + 1
            
            sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
            sorted_phrases = sorted(phrase_freq.items(), key=lambda x: x[1], reverse=True)
            
            top_topics = [word for word, freq in sorted_words[:8] if freq > 2]
            top_phrases = [phrase for phrase, freq in sorted_phrases[:5] if freq > 1]
            
            return top_topics + top_phrases
        except Exception:
            return []
    
    def analyze_content_complexity(self, content: str) -> Dict:
        """Analyze content complexity and characteristics"""
        sentences = sent_tokenize(content)
        words = word_tokenize(content)
        
        avg_sentence_length = len(words) / len(sentences) if sentences else 0
        unique_words = len(set(word.lower() for word in words if word.isalpha()))
        vocabulary_diversity = unique_words / len(words) if words else 0
        
        complex_words = [word for word in words if len(word) > 6]
        technical_indicators = ['analysis', 'methodology', 'implementation', 'framework', 'algorithm']
        technical_score = sum(1 for word in words if word.lower() in technical_indicators)
        
        return {
            'total_sentences': len(sentences),
            'total_words': len(words),
            'avg_sentence_length': round(avg_sentence_length, 2),
            'vocabulary_diversity': round(vocabulary_diversity, 3),
            'complex_words_ratio': round(len(complex_words) / len(words), 3) if words else 0,
            'technical_score': technical_score,
            'estimated_reading_level': 'Basic' if avg_sentence_length < 15 else 'Intermediate' if avg_sentence_length < 25 else 'Advanced'
        }
    
    def create_enhanced_bloom_prompt(self, content: str, settings: Dict, topics_str: str, complexity: Dict) -> str:
        """Create enhanced prompt focused on Bloom's Taxonomy"""
        
        # Map question types to Bloom's levels for better distribution
        bloom_distribution = self.calculate_bloom_distribution(settings)
        
        prompt = f"""You are an expert educational assessment creator specializing in Bloom's Taxonomy. Generate diverse, high-quality questions from the provided content.

CONTENT TO ANALYZE:
{content[:2500]}

CONTENT ANALYSIS:
- Reading Level: {complexity['estimated_reading_level']}
- Key Topics: {topics_str}
- Word Count: {complexity['total_words']}

BLOOM'S TAXONOMY REQUIREMENTS:
Generate questions across ALL levels of Bloom's Taxonomy:

1. REMEMBER (Level 1): {bloom_distribution['remember']} questions
   - Focus: Recall facts, definitions, basic concepts
   - Verbs: define, list, identify, name, state, describe
   - Example: "What is the definition of [concept]?"

2. UNDERSTAND (Level 2): {bloom_distribution['understand']} questions
   - Focus: Explain ideas, summarize, interpret
   - Verbs: explain, summarize, paraphrase, compare, interpret
   - Example: "Explain the main concept of [topic]"

3. APPLY (Level 3): {bloom_distribution['apply']} questions
   - Focus: Use information in new situations
   - Verbs: apply, demonstrate, solve, use, implement
   - Example: "How would you apply [concept] to solve [problem]?"

4. ANALYZE (Level 4): {bloom_distribution['analyze']} questions
   - Focus: Break down information, find relationships
   - Verbs: analyze, compare, examine, categorize
   - Example: "What is the relationship between [concept1] and [concept2]?"

5. EVALUATE (Level 5): {bloom_distribution['evaluate']} questions
   - Focus: Make judgments, critique, assess
   - Verbs: evaluate, judge, critique, defend, justify
   - Example: "Evaluate the effectiveness of [approach]"

6. CREATE (Level 6): {bloom_distribution['create']} questions
   - Focus: Generate new ideas, design solutions
   - Verbs: create, design, formulate, develop, compose
   - Example: "Design a solution that incorporates [concept]"

QUESTION TYPE REQUIREMENTS:
Generate questions using these types: {', '.join([code for code, desc in QUESTION_TYPES if code in settings['question_types']])}

QUESTION TYPE SPECIFICATIONS:
- MCQ: 4 options (A-D), one correct answer, plausible distractors
- FIB: Clear blanks with specific answers, avoid ambiguity
- SAQ: Require 2-3 sentence responses, specific focus
- LAQ: Require detailed explanations, multiple paragraphs
- TRUE_FALSE: Clear statements, avoid absolute terms
- MATCHING: 4-6 pairs, logical connections
- ORDERING: 4-5 items in logical sequence
- CLASSIFICATION: Multiple items, 3-4 categories
- CASE_STUDY: Realistic scenarios requiring analysis

ENHANCED REQUIREMENTS:
- Difficulty levels: {', '.join(settings['difficulty_levels'])}
- Total questions: {settings['num_questions']}
- Include realistic distractors for MCQ
- Provide detailed explanations
- Add prerequisite knowledge
- Include common mistakes
- Suggest follow-up questions

OUTPUT FORMAT (JSON):
{{
  "questions": [
    {{
      "id": "q001",
      "difficulty": "easy|medium|hard",
      "difficulty_score": 1-10,
      "type": "MCQ|FIB|SAQ|LAQ|TRUE_FALSE|MATCHING|ORDERING|CLASSIFICATION|CASE_STUDY",
      "type_description": "Full question type name",
      "cognitive_level": "remember|understand|apply|analyze|evaluate|create",
      "cognitive_description": "Description of cognitive level",
      "question": "Question text",
      "options": ["A", "B", "C", "D"] (for MCQ),
      "correct_answer": "The correct answer",
      "explanation": "Detailed explanation with reasoning",
      "keywords": ["keyword1", "keyword2"],
      "estimated_time": 3,
      "points": 10,
      "prerequisites": ["required knowledge"],
      "common_mistakes": ["typical errors"],
      "follow_up": "Suggested follow-up question",
      "matching_pairs": [{{"term": "definition"}}] (for MATCHING),
      "sequence_items": ["item1", "item2"] (for ORDERING),
      "categories": ["cat1", "cat2"] (for CLASSIFICATION),
      "scenario": "Scenario text" (for CASE_STUDY)
    }}
  ],
  "metadata": {{
    "total_questions": {settings['num_questions']},
    "bloom_distribution": {bloom_distribution},
    "content_topics": ["{topics_str}"],
    "difficulty_distribution": {{"easy": 0, "medium": 0, "hard": 0}}
  }}
}}"""
        
        return prompt
    
    def calculate_bloom_distribution(self, settings: Dict) -> Dict[str, int]:
        """Calculate distribution of questions across Bloom's levels"""
        total_questions = settings['num_questions']
        
        # Default distribution based on educational best practices
        if total_questions <= 5:
            distribution = {'remember': 2, 'understand': 2, 'apply': 1, 'analyze': 0, 'evaluate': 0, 'create': 0}
        elif total_questions <= 10:
            distribution = {'remember': 2, 'understand': 3, 'apply': 2, 'analyze': 2, 'evaluate': 1, 'create': 0}
        elif total_questions <= 20:
            distribution = {'remember': 3, 'understand': 4, 'apply': 4, 'analyze': 4, 'evaluate': 3, 'create': 2}
        else:
            # For larger sets, maintain proportions
            base_dist = [0.2, 0.25, 0.2, 0.15, 0.1, 0.1]  # remember, understand, apply, analyze, evaluate, create
            distribution = {
                'remember': max(1, int(total_questions * base_dist[0])),
                'understand': max(1, int(total_questions * base_dist[1])),
                'apply': max(1, int(total_questions * base_dist[2])),
                'analyze': max(1, int(total_questions * base_dist[3])),
                'evaluate': max(1, int(total_questions * base_dist[4])),
                'create': max(1, int(total_questions * base_dist[5]))
            }
        
        # Adjust to match exact total
        current_total = sum(distribution.values())
        if current_total < total_questions:
            # Add remaining to understand and apply levels
            remaining = total_questions - current_total
            distribution['understand'] += remaining // 2
            distribution['apply'] += remaining - (remaining // 2)
        elif current_total > total_questions:
            # Reduce from create and evaluate levels first
            excess = current_total - total_questions
            for level in ['create', 'evaluate', 'analyze']:
                if excess > 0 and distribution[level] > 1:
                    reduction = min(excess, distribution[level] - 1)
                    distribution[level] -= reduction
                    excess -= reduction
        
        return distribution
    
    def generate_questions_with_provider(self, content: str, settings: Dict, provider_name: str) -> Dict:
        """Generate questions using specified provider via LangChain"""
        if provider_name not in self.model_provider.get_available_models():
            return {"error": f"Provider {provider_name} not available"}
        
        try:
            topics = self.extract_key_topics(content)
            topics_str = ", ".join(topics[:5]) if topics else "general concepts"
            
            complexity = self.analyze_content_complexity(content)
            
            prompt = self.create_enhanced_bloom_prompt(content, settings, topics_str, complexity)
            
            result = self.model_provider.generate_questions(provider_name, prompt, settings)
            
            if 'error' not in result:
                result.setdefault('metadata', {})
                result['metadata'].update({
                    'content_analysis': complexity,
                    'key_topics': topics,
                    'generation_timestamp': datetime.now().isoformat(),
                    'bloom_taxonomy_focus': True
                })
                
                self.analytics_data.append({
                    'provider': provider_name,
                    'timestamp': datetime.now(),
                    'questions_generated': len(result.get('questions', [])),
                    'cost': result.get('usage', {}).get('cost', 0),
                    'content_length': len(content)
                })
            
            return result
            
        except Exception as e:
            return {
                "error": f"Generation failed with {provider_name}: {str(e)}",
                "details": traceback.format_exc()
            }
    
    def ensemble_generation(self, content: str, settings: Dict) -> Dict:
        """Generate questions using multiple providers and merge results"""
        available_models = self.model_provider.get_available_models()
        
        if len(available_models) < 2:
            provider_name = available_models[0] if available_models else None
            if not provider_name:
                return {"error": "No providers available"}
            return self.generate_questions_with_provider(content, settings, provider_name)
        
        results = []
        costs = []
        
        selected_providers = settings.get('selected_providers', available_models[:2])
        
        for provider_name in selected_providers:
            if provider_name in available_models:
                result = self.generate_questions_with_provider(content, settings, provider_name)
                if 'error' not in result:
                    results.append(result)
                    costs.append(result.get('usage', {}).get('cost', 0))
        
        if not results:
            return {"error": "All providers failed to generate questions"}
        
        return self.merge_ensemble_results(results, costs)
    
    def merge_ensemble_results(self, results: List[Dict], costs: List[float]) -> Dict:
        """Intelligently merge results from multiple providers"""
        all_questions = []
        seen_questions = set()
        
        for result in results:
            for question in result.get('questions', []):
                question_key = question['question'].lower().replace(' ', '')[:50]
                if question_key not in seen_questions:
                    seen_questions.add(question_key)
                    all_questions.append(question)
        
        all_questions.sort(key=lambda q: (
            q.get('difficulty_score', 5),
            len(q.get('explanation', '')),
            len(q.get('keywords', []))

        ))
        
        # Limit to requested number of questions
        if len(all_questions) > settings.get('num_questions', 10):
            all_questions = all_questions[:settings['num_questions']]
        
        # Calculate combined metadata
        combined_metadata = {
            'total_questions': len(all_questions),
            'providers_used': len(results),
            'total_cost': sum(costs),
            'generation_timestamp': datetime.now().isoformat(),
            'ensemble_method': 'intelligent_merge',
            'question_distribution': self.analyze_question_distribution(all_questions)
        }
        
        return {
            'questions': all_questions,
            'metadata': combined_metadata,
            'usage': {'cost': sum(costs), 'providers': len(results)}
        }
    
    def analyze_question_distribution(self, questions: List[Dict]) -> Dict:
        """Analyze the distribution of generated questions"""
        distribution = {
            'by_type': {},
            'by_difficulty': {},
            'by_bloom_level': {},
            'avg_time': 0,
            'total_points': 0
        }
        
        for question in questions:
            # Count by type
            q_type = question.get('type', 'unknown')
            distribution['by_type'][q_type] = distribution['by_type'].get(q_type, 0) + 1
            
            # Count by difficulty
            difficulty = question.get('difficulty', 'medium')
            distribution['by_difficulty'][difficulty] = distribution['by_difficulty'].get(difficulty, 0) + 1
            
            # Count by Bloom's level
            bloom_level = question.get('cognitive_level', 'understand')
            distribution['by_bloom_level'][bloom_level] = distribution['by_bloom_level'].get(bloom_level, 0) + 1
            
            # Sum time and points
            distribution['avg_time'] += question.get('estimated_time', 0)
            distribution['total_points'] += question.get('points', 0)
        
        if questions:
            distribution['avg_time'] = round(distribution['avg_time'] / len(questions), 1)
        
        return distribution
    
    def export_questions(self, questions: List[Dict], format_type: str, metadata: Dict = None) -> bytes:
        """Export questions in various formats"""
        if format_type == 'json':
            return self.export_to_json(questions, metadata)
        elif format_type == 'docx':
            return self.export_to_docx(questions, metadata)
        elif format_type == 'xlsx':
            return self.export_to_excel(questions, metadata)
        elif format_type == 'pdf':
            return self.export_to_pdf(questions, metadata)
        elif format_type == 'csv':
            return self.export_to_csv(questions)
        else:
            raise ValueError(f"Unsupported export format: {format_type}")
    
    def export_to_json(self, questions: List[Dict], metadata: Dict = None) -> bytes:
        """Export questions to JSON format"""
        export_data = {
            'questions': questions,
            'metadata': metadata or {},
            'export_timestamp': datetime.now().isoformat(),
            'total_questions': len(questions)
        }
        return json.dumps(export_data, indent=2, ensure_ascii=False).encode('utf-8')
    
    def export_to_docx(self, questions: List[Dict], metadata: Dict = None) -> bytes:
        """Export questions to DOCX format"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Add title
            title = doc.add_heading('Generated Questions', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata section
            if metadata:
                doc.add_heading('Generation Summary', level=1)
                doc.add_paragraph(f"Total Questions: {len(questions)}")
                doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                
                if 'question_distribution' in metadata:
                    dist = metadata['question_distribution']
                    doc.add_paragraph(f"Average Time: {dist.get('avg_time', 0)} minutes")
                    doc.add_paragraph(f"Total Points: {dist.get('total_points', 0)}")
            
            doc.add_page_break()
            
            # Add questions
            for i, question in enumerate(questions, 1):
                doc.add_heading(f'Question {i}', level=2)
                
                # Question metadata
                metadata_p = doc.add_paragraph()
                metadata_p.add_run(f"Type: {question.get('type_description', 'N/A')} | ")
                metadata_p.add_run(f"Difficulty: {question.get('difficulty', 'N/A')} | ")
                metadata_p.add_run(f"Bloom's Level: {question.get('cognitive_level', 'N/A')} | ")
                metadata_p.add_run(f"Time: {question.get('estimated_time', 0)} min | ")
                metadata_p.add_run(f"Points: {question.get('points', 0)}")
                
                # Question text
                doc.add_paragraph(f"Q: {question.get('question', '')}")
                
                # Options for MCQ
                if question.get('options'):
                    for j, option in enumerate(question['options']):
                        doc.add_paragraph(f"{chr(65+j)}. {option}")
                
                # Correct answer
                doc.add_paragraph(f"Answer: {question.get('correct_answer', '')}")
                
                # Explanation
                if question.get('explanation'):
                    doc.add_paragraph(f"Explanation: {question['explanation']}")
                
                # Keywords
                if question.get('keywords'):
                    doc.add_paragraph(f"Keywords: {', '.join(question['keywords'])}")
                
                # Common mistakes
                if question.get('common_mistakes'):
                    doc.add_paragraph(f"Common Mistakes: {'; '.join(question['common_mistakes'])}")
                
                doc.add_paragraph()  # Add space between questions
            
            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            raise ImportError("python-docx package is required for DOCX export")
    
    def export_to_excel(self, questions: List[Dict], metadata: Dict = None) -> bytes:
        """Export questions to Excel format"""
        buffer = io.BytesIO()
        
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            # Questions sheet
            questions_data = []
            for i, q in enumerate(questions, 1):
                row = {
                    'Question_ID': q.get('id', f'q{i:03d}'),
                    'Question_Number': i,
                    'Question_Text': q.get('question', ''),
                    'Question_Type': q.get('type_description', ''),
                    'Difficulty': q.get('difficulty', ''),
                    'Difficulty_Score': q.get('difficulty_score', 0),
                    'Bloom_Level': q.get('cognitive_level', ''),
                    'Correct_Answer': q.get('correct_answer', ''),
                    'Explanation': q.get('explanation', ''),
                    'Keywords': ', '.join(q.get('keywords', [])),
                    'Estimated_Time': q.get('estimated_time', 0),
                    'Points': q.get('points', 0),
                    'Prerequisites': ', '.join(q.get('prerequisites', [])),
                    'Common_Mistakes': '; '.join(q.get('common_mistakes', []))
                }
                
                # Add options for MCQ
                if q.get('options'):
                    for j, option in enumerate(q['options'][:4]):
                        row[f'Option_{chr(65+j)}'] = option
                
                questions_data.append(row)
            
            df_questions = pd.DataFrame(questions_data)
            df_questions.to_excel(writer, sheet_name='Questions', index=False)
            
            # Summary sheet
            if metadata and 'question_distribution' in metadata:
                dist = metadata['question_distribution']
                summary_data = []
                
                # Type distribution
                for q_type, count in dist.get('by_type', {}).items():
                    summary_data.append({'Category': 'Question Type', 'Item': q_type, 'Count': count})
                
                # Difficulty distribution
                for difficulty, count in dist.get('by_difficulty', {}).items():
                    summary_data.append({'Category': 'Difficulty', 'Item': difficulty, 'Count': count})
                
                # Bloom's level distribution
                for bloom_level, count in dist.get('by_bloom_level', {}).items():
                    summary_data.append({'Category': 'Bloom Level', 'Item': bloom_level, 'Count': count})
                
                df_summary = pd.DataFrame(summary_data)
                df_summary.to_excel(writer, sheet_name='Summary', index=False)
        
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_to_csv(self, questions: List[Dict]) -> bytes:
        """Export questions to CSV format"""
        questions_data = []
        for i, q in enumerate(questions, 1):
            row = {
                'Question_ID': q.get('id', f'q{i:03d}'),
                'Question_Number': i,
                'Question_Text': q.get('question', ''),
                'Question_Type': q.get('type_description', ''),
                'Difficulty': q.get('difficulty', ''),
                'Bloom_Level': q.get('cognitive_level', ''),
                'Correct_Answer': q.get('correct_answer', ''),
                'Explanation': q.get('explanation', ''),
                'Keywords': ', '.join(q.get('keywords', [])),
                'Estimated_Time': q.get('estimated_time', 0),
                'Points': q.get('points', 0)
            }
            questions_data.append(row)
        
        df = pd.DataFrame(questions_data)
        return df.to_csv(index=False).encode('utf-8')


class LangChainModelProvider:
    """Unified model provider using LangChain"""
    
    def __init__(self):
        self.models = {}
        self.usage_tracking = {}
        self.encoder = tiktoken.get_encoding("cl100k_base")
    
    def initialize_model(self, provider_name: str, config: Dict) -> bool:
        """Initialize a model provider using LangChain"""
        try:
            if provider_name.lower() == 'openai':
                if config.get('use_chat_model', True):
                    self.models[provider_name] = ChatOpenAI(
                        openai_api_key=config['api_key'],
                        model_name=config.get('model', 'gpt-3.5-turbo'),
                        temperature=config.get('temperature', 0.7),
                        max_tokens=config.get('max_tokens', 3000)
                    )
                else:
                    self.models[provider_name] = OpenAI(
                        openai_api_key=config['api_key'],
                        model_name=config.get('model', 'gpt-3.5-turbo-instruct'),
                        temperature=config.get('temperature', 0.7),
                        max_tokens=config.get('max_tokens', 3000)
                    )
            
            elif provider_name.lower() == 'anthropic':
                self.models[provider_name] = ChatAnthropic(
                    anthropic_api_key=config['api_key'],
                    model=config.get('model', 'claude-3-sonnet-20240229'),
                    temperature=config.get('temperature', 0.7),
                    max_tokens=config.get('max_tokens', 3000)
                )
            
            elif provider_name.lower() == 'deepseek':
                self.models[provider_name] = DeepSeekLLM(
                    api_key=config['api_key'],
                    model=config.get('model', 'deepseek-chat'),
                    temperature=config.get('temperature', 0.7),
                    max_tokens=config.get('max_tokens', 3000)
                )
            
            else:
                # Generic API provider
                self.models[provider_name] = CustomAPILLM(
                    api_key=config['api_key'],
                    model=config.get('model', 'default'),
                    base_url=config.get('base_url', ''),
                    provider_name=provider_name,
                    temperature=config.get('temperature', 0.7),
                    max_tokens=config.get('max_tokens', 3000)
                )
            
            self.usage_tracking[provider_name] = {'requests': 0, 'tokens': 0, 'cost': 0.0}
            return True
            
        except Exception as e:
            st.error(f"Failed to initialize {provider_name}: {str(e)}")
            return False
    
    def get_available_models(self) -> List[str]:
        """Get list of available model providers"""
        return list(self.models.keys())
    
    def count_tokens(self, text: str) -> int:
        """Count tokens in text"""
        try:
            return len(self.encoder.encode(text))
        except:
            return len(text.split()) * 1.3  # Rough approximation
    
    def generate_questions(self, provider_name: str, prompt: str, settings: Dict) -> Dict:
        """Generate questions using specified provider"""
        if provider_name not in self.models:
            return {"error": f"Provider {provider_name} not initialized"}
        
        try:
            model = self.models[provider_name]
            
            # Track usage
            self.usage_tracking[provider_name]['requests'] += 1
            prompt_tokens = self.count_tokens(prompt)
            self.usage_tracking[provider_name]['tokens'] += prompt_tokens
            
            # Generate response
            if isinstance(model, (ChatOpenAI, ChatAnthropic)):
                messages = [
                    SystemMessage(content="You are an expert educational assessment creator specializing in Bloom's Taxonomy. Respond only with valid JSON."),
                    HumanMessage(content=prompt)
                ]
                response = model(messages)
                response_text = response.content
            else:
                response_text = model(prompt)
            
            # Track response tokens
            response_tokens = self.count_tokens(response_text)
            self.usage_tracking[provider_name]['tokens'] += response_tokens
            
            # Parse JSON response
            try:
                # Clean response text
                response_text = response_text.strip()
                if response_text.startswith('```json'):
                    response_text = response_text[7:]
                if response_text.endswith('```'):
                    response_text = response_text[:-3]
                
                result = json.loads(response_text)
                
                # Validate and enhance questions
                if 'questions' in result:
                    enhanced_questions = []
                    for i, question in enumerate(result['questions']):
                        enhanced_q = self.enhance_question(question, i + 1)
                        enhanced_questions.append(enhanced_q)
                    result['questions'] = enhanced_questions
                
                # Add usage information
                cost = self.estimate_cost(provider_name, prompt_tokens + response_tokens)
                self.usage_tracking[provider_name]['cost'] += cost
                
                result['usage'] = {
                    'prompt_tokens': prompt_tokens,
                    'response_tokens': response_tokens,
                    'total_tokens': prompt_tokens + response_tokens,
                    'cost': cost,
                    'provider': provider_name
                }
                
                return result
                
            except json.JSONDecodeError as e:
                return {
                    "error": f"Failed to parse JSON response from {provider_name}",
                    "details": str(e),
                    "raw_response": response_text[:500]
                }
        
        except Exception as e:
            return {
                "error": f"Generation failed with {provider_name}",
                "details": str(e)
            }
    
    def enhance_question(self, question: Dict, question_num: int) -> Dict:
        """Enhance question with additional metadata and validation"""
        enhanced = question.copy()
        
        # Ensure required fields
        enhanced.setdefault('id', f'q{question_num:03d}')
        enhanced.setdefault('difficulty', 'medium')
        enhanced.setdefault('difficulty_score', 5)
        enhanced.setdefault('type', 'SAQ')
        enhanced.setdefault('cognitive_level', 'understand')
        enhanced.setdefault('estimated_time', 3)
        enhanced.setdefault('points', 10)
        enhanced.setdefault('keywords', [])
        enhanced.setdefault('prerequisites', [])
        enhanced.setdefault('common_mistakes', [])
        
        # Validate cognitive level
        if enhanced['cognitive_level'] not in BLOOMS_TAXONOMY:
            enhanced['cognitive_level'] = 'understand'
        
        # Add cognitive description
        enhanced['cognitive_description'] = BLOOMS_TAXONOMY[enhanced['cognitive_level']]['description']
        
        # Validate question type
        valid_types = [code for code, desc in QUESTION_TYPES]
        if enhanced['type'] not in valid_types:
            enhanced['type'] = 'SAQ'
        
        # Add type description
        for code, desc in QUESTION_TYPES:
            if enhanced['type'] == code:
                enhanced['type_description'] = desc
                break
        
        # Validate difficulty score
        if not isinstance(enhanced.get('difficulty_score'), int) or not (1 <= enhanced['difficulty_score'] <= 10):
            difficulty_map = {'easy': 3, 'medium': 5, 'hard': 8}
            enhanced['difficulty_score'] = difficulty_map.get(enhanced.get('difficulty', 'medium'), 5)
        
        return enhanced
    
    def estimate_cost(self, provider_name: str, total_tokens: int) -> float:
        """Estimate cost based on provider and token count"""
        # Cost estimates per 1K tokens (input + output combined for simplicity)
        cost_rates = {
            'openai': 0.002,  # GPT-3.5-turbo average
            'anthropic': 0.008,  # Claude-3-sonnet average
            'deepseek': 0.0014,  # DeepSeek-chat
            'gemini': 0.001,  # Gemini Pro
            'groq': 0.0002,  # Groq Mixtral
        }
        
        rate = cost_rates.get(provider_name.lower(), 0.002)
        return (total_tokens / 1000) * rate
    
    def get_usage_summary(self) -> Dict:
        """Get usage summary across all providers"""
        return self.usage_tracking.copy()


# Streamlit App Main Function
def main():
    st.set_page_config(
        page_title="🎓 Advanced Question Generator",
        page_icon="🎓",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Initialize session state
    if 'generator' not in st.session_state:
        st.session_state.generator = BloomsTaxonomyQuestionGenerator()
    
    if 'generated_questions' not in st.session_state:
        st.session_state.generated_questions = []
    
    if 'generation_metadata' not in st.session_state:
        st.session_state.generation_metadata = {}
    
    # App header
    st.title("🎓 Advanced Question Generator")
    st.markdown("*Generate high-quality educational questions using Bloom's Taxonomy and multiple AI providers*")
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # Model Provider Configuration
        st.subheader("🤖 AI Model Providers")
        
        provider_configs = {}
        
        # OpenAI Configuration
        with st.expander("OpenAI Configuration"):
            openai_key = st.text_input("OpenAI API Key", type="password", key="openai_key")
            openai_model = st.selectbox(
                "Model", 
                ["gpt-4", "gpt-4-turbo", "gpt-3.5-turbo", "gpt-3.5-turbo-16k"],
                key="openai_model"
            )
            if openai_key:
                provider_configs['openai'] = {
                    'api_key': openai_key,
                    'model': openai_model,
                    'use_chat_model': True
                }
        
        # Anthropic Configuration
        with st.expander("Anthropic Configuration"):
            anthropic_key = st.text_input("Anthropic API Key", type="password", key="anthropic_key")
            anthropic_model = st.selectbox(
                "Model",
                ["claude-3-opus-20240229", "claude-3-sonnet-20240229", "claude-3-haiku-20240307"],
                key="anthropic_model"
            )
            if anthropic_key:
                provider_configs['anthropic'] = {
                    'api_key': anthropic_key,
                    'model': anthropic_model
                }
        
        # DeepSeek Configuration
        with st.expander("DeepSeek Configuration"):
            deepseek_key = st.text_input("DeepSeek API Key", type="password", key="deepseek_key")
            if deepseek_key:
                provider_configs['deepseek'] = {
                    'api_key': deepseek_key,
                    'model': 'deepseek-chat'
                }
        
        # Initialize providers
        if provider_configs:
            if st.button("🔄 Initialize Providers"):
                with st.spinner("Initializing AI providers..."):
                    errors = st.session_state.generator.initialize_providers(provider_configs)
                    if errors:
                        for provider, error in errors.items():
                            st.error(f"{provider}: {error}")
                    else:
                        st.success("All providers initialized successfully!")
        
        # Generation Settings
        st.subheader("📝 Generation Settings")
        
        num_questions = st.slider("Number of Questions", 1, 50, 10)
        
        difficulty_levels = st.multiselect(
            "Difficulty Levels",
            ["easy", "medium", "hard"],
            default=["medium", "hard"]
        )
        
        question_types = st.multiselect(
            "Question Types",
            [code for code, desc in QUESTION_TYPES],
            default=["MCQ", "SAQ", "LAQ"]
        )
        
        # Advanced settings
        with st.expander("🔧 Advanced Settings"):
            use_ensemble = st.checkbox("Use Ensemble Generation", value=False)
            max_chunk_size = st.slider("Max Content Chunk Size (tokens)", 1000, 4000, 2000)
            include_prerequisites = st.checkbox("Include Prerequisites", value=True)
            include_common_mistakes = st.checkbox("Include Common Mistakes", value=True)
    
    # Main content area
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("📄 Content Input")
        
        # Content input methods
        input_method = st.radio(
            "Content Input Method",
            ["Upload File", "Paste Text", "Enter URL"],
            horizontal=True
        )
        
        content = ""
        
        if input_method == "Upload File":
            uploaded_file = st.file_uploader(
                "Upload your content file",
                type=['pdf', 'docx', 'txt', 'csv', 'xlsx'],
                help="Supported formats: PDF, DOCX, TXT, CSV, XLSX"
            )
            
            if uploaded_file:
                with st.spinner("Extracting content..."):
                    content = st.session_state.generator.extract_content(uploaded_file)
                
                if content:
                    st.success(f"Content extracted successfully! ({len(content)} characters)")
                    with st.expander("Preview Content"):
                        st.text_area("Extracted Content", content[:1000] + "..." if len(content) > 1000 else content, height=200)
        
        elif input_method == "Paste Text":
            content = st.text_area(
                "Paste your content here",
                height=300,
                placeholder="Paste your educational content here..."
            )
        
        elif input_method == "Enter URL":
            url = st.text_input("Enter URL", placeholder="https://example.com/article")
            if url and st.button("Fetch Content"):
                st.info("URL content fetching would be implemented here")
        
        # Content analysis
        if content:
            with st.expander("📊 Content Analysis"):
                complexity = st.session_state.generator.analyze_content_complexity(content)
                topics = st.session_state.generator.extract_key_topics(content)
                
                col_a, col_b, col_c = st.columns(3)
                with col_a:
                    st.metric("Total Words", complexity['total_words'])
                    st.metric("Sentences", complexity['total_sentences'])
                
                with col_b:
                    st.metric("Avg Sentence Length", complexity['avg_sentence_length'])
                    st.metric("Reading Level", complexity['estimated_reading_level'])
                
                with col_c:
                    st.metric("Vocabulary Diversity", complexity['vocabulary_diversity'])
                    st.metric("Technical Score", complexity['technical_score'])
                
                if topics:
                    st.subheader("Key Topics Identified")
                    st.write(", ".join(topics[:10]))
    
    with col2:
        st.header("🎯 Generation Control")
        
        available_models = st.session_state.generator.model_provider.get_available_models()
        
        if not available_models:
            st.warning("⚠️ No AI providers configured. Please configure at least one provider in the sidebar.")
        else:
            st.success(f"✅ {len(available_models)} provider(s) available: {', '.join(available_models)}")
            
            if use_ensemble and len(available_models) > 1:
                selected_providers = st.multiselect(
                    "Select Providers for Ensemble",
                    available_models,
                    default=available_models[:2]
                )
            else:
                selected_provider = st.selectbox("Select Provider", available_models)
                selected_providers = [selected_provider]
        
        # Generate button
        if content and available_models:
            if st.button("🚀 Generate Questions", type="primary"):
                settings = {
                    'num_questions': num_questions,
                    'difficulty_levels': difficulty_levels,
                    'question_types': question_types,
                    'selected_providers': selected_providers,
                    'max_chunk_size': max_chunk_size,
                    'include_prerequisites': include_prerequisites,
                    'include_common_mistakes': include_common_mistakes
                }
                
                with st.spinner("Generating questions..."):
                    if use_ensemble and len(selected_providers) > 1:
                        result = st.session_state.generator.ensemble_generation(content, settings)
                    else:
                        result = st.session_state.generator.generate_questions_with_provider(
                            content, settings, selected_providers[0]
                        )
                
                if 'error' in result:
                    st.error(f"Generation failed: {result['error']}")
                    if 'details' in result:
                        st.error(f"Details: {result['details']}")
                else:
                    st.session_state.generated_questions = result.get('questions', [])
                    st.session_state.generation_metadata = result.get('metadata', {})
                    st.success(f"✅ Generated {len(st.session_state.generated_questions)} questions!")
                    st.rerun()
    
    # Display generated questions
    if st.session_state.generated_questions:
        st.header("📋 Generated Questions")
        
        # Summary metrics
        questions = st.session_state.generated_questions
        metadata = st.session_state.generation_metadata
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Questions", len(questions))
        with col2:
            avg_time = sum(q.get('estimated_time', 0) for q in questions) / len(questions)
            st.metric("Avg Time (min)", f"{avg_time:.1f}")
        with col3:
            total_points = sum(q.get('points', 0) for q in questions)
            st.metric("Total Points", total_points)
        with col4:
            if 'usage' in metadata:
                st.metric("Generation Cost", f"${metadata['usage'].get('cost', 0):.4f}")
        
        # Question distribution charts
        if 'question_distribution' in metadata:
            dist = metadata['question_distribution']
            
            col1, col2 = st.columns(2)
            
            with col1:
                if dist.get('by_bloom_level'):
                    fig_bloom = px.pie(
                        values=list(dist['by_bloom_level'].values()),
                        names=list(dist['by_bloom_level'].keys()),
                        title="Distribution by Bloom's Taxonomy Level"
                    )
                    st.plotly_chart(fig_bloom, use_container_width=True)
            
            with col2:
                if dist.get('by_difficulty'):
                    fig_diff = px.bar(
                        x=list(dist['by_difficulty'].keys()),
                        y=list(dist['by_difficulty'].values()),
                        title="Distribution by Difficulty Level"
                    )
                    st.plotly_chart(fig_diff, use_container_width=True)
        
        # Question display and editing
        st.subheader("📝 Questions")
        
        # Filter and sort options
        col1, col2, col3 = st.columns(3)
        with col1:
            filter_type = st.selectbox("Filter by Type", ["All"] + [q.get('type', 'Unknown') for q in questions])
        with col2:
            filter_difficulty = st.selectbox("Filter by Difficulty", ["All", "easy", "medium", "hard"])
        with col3:
            sort_by = st.selectbox("Sort by", ["Question Number", "Difficulty", "Bloom Level", "Type"])
        
                # Apply filters
        filtered_questions = questions.copy()
        if filter_type != "All":
            filtered_questions = [q for q in filtered_questions if q.get('type', 'Unknown') == filter_type]
        if filter_difficulty != "All":
            filtered_questions = [q for q in filtered_questions if q.get('difficulty', 'medium') == filter_difficulty]

        # Apply sorting
        if sort_by == "Difficulty":
            filtered_questions.sort(key=lambda q: q.get('difficulty_score', 5))
        elif sort_by == "Bloom Level":
            filtered_questions.sort(key=lambda q: BLOOMS_TAXONOMY.get(q.get('cognitive_level', 'understand'), {}).get('level', 2))
        elif sort_by == "Type":
            filtered_questions.sort(key=lambda q: q.get('type', ''))
        else:
            filtered_questions.sort(key=lambda q: int(q.get('id', 'q001').lstrip('q')))

        # Display filtered questions
        for idx, question in enumerate(filtered_questions, 1):
            with st.expander(f"Question {idx}: {question.get('question', '')[:60]}..."):
                st.markdown(f"**Type:** {question.get('type_description', '')}")
                st.markdown(f"**Difficulty:** {question.get('difficulty', '')} (Score: {question.get('difficulty_score', 5)})")
                st.markdown(f"**Bloom Level:** {question.get('cognitive_level', '')}")
                st.markdown(f"**Estimated Time:** {question.get('estimated_time', 0)} min")
                st.markdown(f"**Points:** {question.get('points', 0)}")

                st.markdown(f"**Question:** {question.get('question', '')}")

                if question.get('options'):
                    for i, opt in enumerate(question['options']):
                        st.markdown(f"- {chr(65 + i)}. {opt}")

                st.markdown(f"**Answer:** {question.get('correct_answer', '')}")
                st.markdown(f"**Explanation:** {question.get('explanation', '')}")
                if question.get('keywords'):
                    st.markdown(f"**Keywords:** {', '.join(question.get('keywords'))}")
                if question.get('prerequisites'):
                    st.markdown(f"**Prerequisites:** {', '.join(question.get('prerequisites'))}")
                if question.get('common_mistakes'):
                    st.markdown(f"**Common Mistakes:** {', '.join(question.get('common_mistakes'))}")
                if question.get('follow_up'):
                    st.markdown(f"**Follow-up:** {question.get('follow_up')}")

        # Export options
        st.subheader("📤 Export Questions")
        export_format = st.selectbox("Select Export Format", ["json", "docx", "xlsx", "csv"])
        if st.button("Download Export File"):
            export_data = st.session_state.generator.export_questions(filtered_questions, export_format, metadata)
            st.download_button(
                label=f"Download {export_format.upper()} file",
                data=export_data,
                file_name=f"generated_questions.{export_format}",
                mime="application/octet-stream"
            )

        # Display usage summary
        st.subheader("📈 Usage Summary")
        usage_summary = st.session_state.generator.model_provider.get_usage_summary()
        for provider, usage in usage_summary.items():
            st.markdown(f"**{provider.capitalize()}**")
            st.markdown(f"- Requests: {usage['requests']}")
            st.markdown(f"- Tokens Used: {usage['tokens']}")
            st.markdown(f"- Cost: ${usage['cost']:.4f}")
if __name__ == "__main__":
    main()
