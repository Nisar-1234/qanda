import streamlit as st
import os
import io
import json
import pandas as pd
import docx
import PyPDF2
import requests
import traceback
import re
from datetime import datetime
import hashlib
from typing import Dict, List, Any, Tuple, Optional
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
import tiktoken
import plotly.express as px
import plotly.graph_objects as go

# LangChain imports
from langchain.llms import OpenAI
from langchain.chat_models import ChatOpenAI, ChatAnthropic
from langchain.llms.base import LLM
from langchain.schema import BaseMessage, HumanMessage, SystemMessage
from langchain.callbacks.manager import CallbackManagerForLLMRun
from langchain.schema.output import LLMResult
from langchain.prompts import PromptTemplate, ChatPromptTemplate
from langchain.output_parsers import PydanticOutputParser, OutputFixingParser
from langchain.schema import OutputParserException
from pydantic import BaseModel, Field
import yaml

# Download required NLTK data (run once)
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
except LookupError:
    with st.spinner("Downloading language models..."):
        nltk.download('punkt', quiet=True)
        nltk.download('stopwords', quiet=True)

# Pydantic models for structured output
class QuestionModel(BaseModel):
    id: str = Field(description="Unique identifier for the question")
    difficulty: str = Field(description="Difficulty level: easy, medium, or hard")
    difficulty_score: int = Field(description="Difficulty score from 1-10")
    type: str = Field(description="Question type: multiple_choice, true_false, short_answer, essay, fill_blank")
    cognitive_level: str = Field(description="Bloom's taxonomy level")
    question: str = Field(description="The question text")
    options: List[str] = Field(default=[], description="Multiple choice options if applicable")
    correct_answer: str = Field(description="The correct answer")
    explanation: str = Field(description="Detailed explanation")
    keywords: List[str] = Field(default=[], description="Key terms and concepts")
    estimated_time: int = Field(description="Estimated time in minutes")
    points: int = Field(description="Point value for the question")
    prerequisites: List[str] = Field(default=[], description="Required prerequisite knowledge")
    common_mistakes: List[str] = Field(default=[], description="Common student errors")
    follow_up: str = Field(default="", description="Follow-up question suggestion")

class QuestionSetModel(BaseModel):
    questions: List[QuestionModel] = Field(description="List of generated questions")
    metadata: Dict[str, Any] = Field(description="Generation metadata")

# Custom LangChain LLM for DeepSeek
class DeepSeekLLM(LLM):
    api_key: str
    model: str = "deepseek-chat"
    base_url: str = "https://api.deepseek.com/v1"
    temperature: float = 0.7
    max_tokens: int = 3000
    
    @property
    def _llm_type(self) -> str:
        return "deepseek"
    
    def _call(
        self,
        prompt: str,
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs: Any,
    ) -> str:
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": self.model,
            "messages": [
                {"role": "system", "content": "You are an expert educational assessment creator. Respond only with valid JSON."},
                {"role": "user", "content": prompt}
            ],
            "temperature": self.temperature,
            "max_tokens": self.max_tokens
        }
        
        response = requests.post(f"{self.base_url}/chat/completions", headers=headers, json=data)
        response.raise_for_status()
        
        response_data = response.json()
        return response_data['choices'][0]['message']['content']

# Custom LangChain LLM for other providers
class CustomAPILLM(LLM):
    api_key: str
    model: str
    base_url: str
    provider_name: str
    temperature: float = 0.7
    max_tokens: int = 3000
    
    @property
    def _llm_type(self) -> str:
        return self.provider_name.lower()
    
    def _call(
        self,
        prompt: str,
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs: Any,
    ) -> str:
        # Generic implementation for API-based LLMs
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": self.model,
            "messages": [
                {"role": "system", "content": "You are an expert educational assessment creator. Respond only with valid JSON."},
                {"role": "user", "content": prompt}
            ],
            "temperature": self.temperature,
            "max_tokens": self.max_tokens
        }
        
        response = requests.post(f"{self.base_url}/chat/completions", headers=headers, json=data)
        response.raise_for_status()
        
        response_data = response.json()
        return response_data['choices'][0]['message']['content']

class LangChainModelProvider:
    """LangChain-based model provider for unified LLM interface"""
    
    def __init__(self):
        self.models = {}
        self.cost_tracking = {}
        self.encoding = tiktoken.get_encoding("cl100k_base")
        
        # Cost per 1K tokens for different models
        self.model_costs = {
            "gpt-4o": 0.15,
            "gpt-4o-mini": 0.002,
            "gpt-4-turbo": 0.10,
            "gpt-3.5-turbo": 0.002,
            "claude-3-opus": 0.25,
            "claude-3-sonnet": 0.125,
            "claude-3-haiku": 0.025,
            "deepseek-chat": 0.002
        }
    
    def initialize_model(self, provider: str, config: Dict) -> bool:
        """Initialize a model provider using LangChain"""
        try:
            if provider == "openai":
                if config.get('chat_model', True):
                    model = ChatOpenAI(
                        openai_api_key=config['api_key'],
                        model_name=config.get('model', 'gpt-4o-mini'),
                        temperature=config.get('temperature', 0.7),
                        max_tokens=config.get('max_tokens', 3000)
                    )
                else:
                    model = OpenAI(
                        openai_api_key=config['api_key'],
                        model_name=config.get('model', 'gpt-3.5-turbo-instruct'),
                        temperature=config.get('temperature', 0.7),
                        max_tokens=config.get('max_tokens', 3000)
                    )
                
                self.models[provider] = {
                    'model': model,
                    'name': f"OpenAI ({config.get('model', 'gpt-4o-mini')})",
                    'cost_per_1k': self.model_costs.get(config.get('model', 'gpt-4o-mini'), 0.002),
                    'type': 'chat' if config.get('chat_model', True) else 'completion'
                }
                return True
                
            elif provider == "claude":
                model = ChatAnthropic(
                    anthropic_api_key=config['api_key'],
                    model=config.get('model', 'claude-3-haiku-20240307'),
                    temperature=config.get('temperature', 0.7),
                    max_tokens=config.get('max_tokens', 3000)
                )
                
                self.models[provider] = {
                    'model': model,
                    'name': f"Claude ({config.get('model', 'claude-3-haiku')})",
                    'cost_per_1k': self.model_costs.get(config.get('model', 'claude-3-haiku'), 0.025),
                    'type': 'chat'
                }
                return True
                
            elif provider == "deepseek":
                model = DeepSeekLLM(
                    api_key=config['api_key'],
                    model=config.get('model', 'deepseek-chat'),
                    temperature=config.get('temperature', 0.7),
                    max_tokens=config.get('max_tokens', 3000)
                )
                
                self.models[provider] = {
                    'model': model,
                    'name': f"DeepSeek ({config.get('model', 'deepseek-chat')})",
                    'cost_per_1k': self.model_costs.get(config.get('model', 'deepseek-chat'), 0.002),
                    'type': 'completion'
                }
                return True
                
            elif provider == "custom":
                model = CustomAPILLM(
                    api_key=config['api_key'],
                    model=config.get('model', 'custom-model'),
                    base_url=config['base_url'],
                    provider_name=config.get('provider_name', 'Custom'),
                    temperature=config.get('temperature', 0.7),
                    max_tokens=config.get('max_tokens', 3000)
                )
                
                self.models[provider] = {
                    'model': model,
                    'name': f"{config.get('provider_name', 'Custom')} ({config.get('model', 'custom-model')})",
                    'cost_per_1k': config.get('cost_per_1k', 0.01),
                    'type': 'completion'
                }
                return True
                
        except Exception as e:
            st.error(f"Failed to initialize {provider}: {str(e)}")
            return False
        
        return False
    
    def count_tokens(self, text: str) -> int:
        """Count tokens in text"""
        return len(self.encoding.encode(text))
    
    def estimate_cost(self, provider: str, input_tokens: int, output_tokens: int = 1000) -> float:
        """Estimate cost for a provider"""
        if provider in self.models:
            cost_per_1k = self.models[provider]['cost_per_1k']
            return (input_tokens + output_tokens) / 1000 * cost_per_1k
        return 0.0
    
    def generate_questions(self, provider: str, prompt: str, settings: Dict) -> Dict:
        """Generate questions using specified provider"""
        if provider not in self.models:
            return {"error": f"Provider {provider} not available"}
        
        model_info = self.models[provider]
        model = model_info['model']
        
        try:
            # Create structured output parser
            parser = PydanticOutputParser(pydantic_object=QuestionSetModel)
            format_instructions = parser.get_format_instructions()
            
            # Add format instructions to prompt
            enhanced_prompt = f"{prompt}\n\n{format_instructions}"
            
            # Track input tokens
            input_tokens = self.count_tokens(enhanced_prompt)
            
            # Generate based on model type
            if model_info['type'] == 'chat':
                messages = [
                    SystemMessage(content="You are an expert educational assessment creator."),
                    HumanMessage(content=enhanced_prompt)
                ]
                response = model(messages)
                content = response.content
            else:
                response = model(enhanced_prompt)
                content = response
            
            # Parse the response
            try:
                # Try to extract JSON if wrapped in code blocks
                if "```json" in content:
                    content = content.split("```json")[1].split("```")[0]
                elif "```" in content:
                    content = content.split("```")[1].split("```")[0]
                
                parsed_result = parser.parse(content)
                result = parsed_result.dict()
                
            except OutputParserException:
                # Fallback: try to parse as regular JSON
                try:
                    result = json.loads(content)
                except json.JSONDecodeError:
                    return {"error": f"Failed to parse response from {provider}"}
            
            # Add metadata
            output_tokens = self.count_tokens(content)
            cost = self.estimate_cost(provider, input_tokens, output_tokens)
            
            result['provider'] = model_info['name']
            result['usage'] = {
                'input_tokens': input_tokens,
                'output_tokens': output_tokens,
                'cost': cost
            }
            
            # Track usage
            if provider not in self.cost_tracking:
                self.cost_tracking[provider] = {'total_cost': 0, 'total_tokens': 0}
            
            self.cost_tracking[provider]['total_cost'] += cost
            self.cost_tracking[provider]['total_tokens'] += input_tokens + output_tokens
            
            return result
            
        except Exception as e:
            return {
                "error": f"Generation failed with {provider}: {str(e)}",
                "details": traceback.format_exc()
            }
    
    def get_available_models(self) -> List[str]:
        """Get list of available model providers"""
        return list(self.models.keys())
    
    def get_model_info(self, provider: str) -> Dict:
        """Get information about a specific model"""
        return self.models.get(provider, {})

class AdvancedQuestionGenerator:
    def __init__(self):
        self.model_provider = LangChainModelProvider()
        self.max_tokens = 3000
        self.question_history = []
        self.analytics_data = []
        
        # Enhanced question templates with more variety
        self.question_templates = {
            'conceptual': [
                "What is the main concept of {topic}?",
                "Explain the significance of {topic}.",
                "How does {topic} relate to {context}?",
                "Define {topic} in your own words.",
                "What are the key characteristics of {topic}?"
            ],
            'analytical': [
                "Analyze the relationship between {concept1} and {concept2}.",
                "What are the implications of {topic}?",
                "Compare and contrast {topic} with similar concepts.",
                "What factors contribute to {topic}?",
                "Evaluate the effectiveness of {topic}."
            ],
            'application': [
                "How would you apply {concept} in a real-world scenario?",
                "What problems does {topic} solve?",
                "Demonstrate the practical use of {concept}.",
                "Create an example that illustrates {topic}.",
                "How might {topic} be used in {context}?"
            ],
            'synthesis': [
                "Combine concepts from {topic1} and {topic2} to create a new solution.",
                "How would you design a system that incorporates {topic}?",
                "What new insights emerge when connecting {concept1} with {concept2}?"
            ]
        }
    
    def initialize_providers(self, provider_configs: Dict) -> Dict[str, str]:
        """Initialize AI model providers using LangChain"""
        errors = {}
        
        for provider, config in provider_configs.items():
            if config.get('api_key'):
                success = self.model_provider.initialize_model(provider, config)
                if not success:
                    errors[provider] = "Failed to initialize"
        
        return errors
    
    def count_tokens(self, text: str) -> int:
        """Count tokens in text"""
        return self.model_provider.count_tokens(text)
    
    def extract_text_from_pdf(self, file) -> str:
        """Enhanced PDF extraction with better error handling"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            total_pages = len(pdf_reader.pages)
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for page_num in range(total_pages):
                try:
                    page_text = pdf_reader.pages[page_num].extract_text()
                    # Enhanced text cleaning
                    page_text = re.sub(r'\n+', '\n', page_text)
                    page_text = re.sub(r'\s+', ' ', page_text)
                    page_text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\"\'\/]', '', page_text)
                    text += page_text + "\n"
                    
                    progress = (page_num + 1) / total_pages
                    progress_bar.progress(progress)
                    status_text.text(f"Processing page {page_num + 1} of {total_pages}")
                    
                except Exception as e:
                    st.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                    continue
            
            progress_bar.empty()
            status_text.empty()
            return text.strip()
            
        except Exception as e:
            st.error(f"Error reading PDF: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, file) -> str:
        """Enhanced DOCX extraction with tables and formatting"""
        try:
            doc = docx.Document(file)
            text = ""
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text.strip() + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            return text.strip()
        except Exception as e:
            st.error(f"Error reading DOCX: {str(e)}")
            return ""
    
    def extract_content(self, uploaded_file) -> str:
        """Extract content from various file types"""
        file_type = uploaded_file.name.split('.')[-1].lower()
        
        try:
            if file_type == 'pdf':
                return self.extract_text_from_pdf(uploaded_file)
            elif file_type in ['docx', 'doc']:
                return self.extract_text_from_docx(uploaded_file)
            elif file_type in ['xlsx', 'xls']:
                df = pd.read_excel(uploaded_file)
                return df.to_string(index=False)
            elif file_type == 'txt':
                return uploaded_file.getvalue().decode('utf-8')
            elif file_type == 'csv':
                df = pd.read_csv(uploaded_file)
                return df.to_string(index=False)
            else:
                return f"Unsupported file type: {file_type}"
        except Exception as e:
            return f"Error extracting content: {str(e)}"
    
    def intelligent_content_chunking(self, content: str, max_chunk_tokens: int = 2000) -> List[str]:
        """Intelligently chunk content by sentences and paragraphs"""
        paragraphs = content.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
                
            test_chunk = current_chunk + "\n\n" + paragraph if current_chunk else paragraph
            
            if self.count_tokens(test_chunk) <= max_chunk_tokens:
                current_chunk = test_chunk
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                
                if self.count_tokens(paragraph) > max_chunk_tokens:
                    sentences = sent_tokenize(paragraph)
                    sentence_chunk = ""
                    
                    for sentence in sentences:
                        test_sentence_chunk = sentence_chunk + " " + sentence if sentence_chunk else sentence
                        
                        if self.count_tokens(test_sentence_chunk) <= max_chunk_tokens:
                            sentence_chunk = test_sentence_chunk
                        else:
                            if sentence_chunk:
                                chunks.append(sentence_chunk)
                            sentence_chunk = sentence
                    
                    if sentence_chunk:
                        current_chunk = sentence_chunk
                else:
                    current_chunk = paragraph
        
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def extract_key_topics(self, content: str) -> List[str]:
        """Extract key topics using enhanced NLP techniques"""
        try:
            words = word_tokenize(content.lower())
            stop_words = set(stopwords.words('english'))
            
            # Enhanced filtering
            filtered_words = [
                word for word in words 
                if word.isalpha() and word not in stop_words and len(word) > 3
                and not word.isdigit()
            ]
            
            # Get word frequency
            word_freq = {}
            for word in filtered_words:
                word_freq[word] = word_freq.get(word, 0) + 1
            
            # Extract multi-word phrases
            phrases = []
            words_list = content.split()
            for i in range(len(words_list) - 1):
                phrase = f"{words_list[i]} {words_list[i+1]}".lower()
                if len(phrase) > 6 and phrase.count(' ') == 1:
                    phrases.append(phrase)
            
            phrase_freq = {}
            for phrase in phrases:
                phrase_freq[phrase] = phrase_freq.get(phrase, 0) + 1
            
            # Combine single words and phrases
            sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
            sorted_phrases = sorted(phrase_freq.items(), key=lambda x: x[1], reverse=True)
            
            top_topics = [word for word, freq in sorted_words[:8] if freq > 2]
            top_phrases = [phrase for phrase, freq in sorted_phrases[:5] if freq > 1]
            
            return top_topics + top_phrases
        except Exception:
            return []
    
    def analyze_content_complexity(self, content: str) -> Dict:
        """Analyze content complexity and characteristics"""
        sentences = sent_tokenize(content)
        words = word_tokenize(content)
        
        # Basic metrics
        avg_sentence_length = len(words) / len(sentences) if sentences else 0
        unique_words = len(set(word.lower() for word in words if word.isalpha()))
        vocabulary_diversity = unique_words / len(words) if words else 0
        
        # Complexity indicators
        complex_words = [word for word in words if len(word) > 6]
        technical_indicators = ['analysis', 'methodology', 'implementation', 'framework', 'algorithm']
        technical_score = sum(1 for word in words if word.lower() in technical_indicators)
        
        return {
            'total_sentences': len(sentences),
            'total_words': len(words),
            'avg_sentence_length': round(avg_sentence_length, 2),
            'vocabulary_diversity': round(vocabulary_diversity, 3),
            'complex_words_ratio': round(len(complex_words) / len(words), 3) if words else 0,
            'technical_score': technical_score,
            'estimated_reading_level': 'Basic' if avg_sentence_length < 15 else 'Intermediate' if avg_sentence_length < 25 else 'Advanced'
        }
    
    def generate_questions_with_provider(self, content: str, settings: Dict, provider_name: str) -> Dict:
        """Generate questions using specified provider via LangChain"""
        if provider_name not in self.model_provider.get_available_models():
            return {"error": f"Provider {provider_name} not available"}
        
        try:
            # Extract key topics for this chunk
            topics = self.extract_key_topics(content)
            topics_str = ", ".join(topics[:5]) if topics else "general concepts"
            
            # Analyze content complexity
            complexity = self.analyze_content_complexity(content)
            
            # Create enhanced prompt
            prompt = self.create_enhanced_prompt(content, settings, topics_str, complexity)
            
            # Generate questions using LangChain
            result = self.model_provider.generate_questions(provider_name, prompt, settings)
            
            if 'error' not in result:
                # Add metadata
                result.setdefault('metadata', {})
                result['metadata'].update({
                    'content_analysis': complexity,
                    'key_topics': topics,
                    'generation_timestamp': datetime.now().isoformat()
                })
                
                # Store for analytics
                self.analytics_data.append({
                    'provider': provider_name,
                    'timestamp': datetime.now(),
                    'questions_generated': len(result.get('questions', [])),
                    'cost': result.get('usage', {}).get('cost', 0),
                    'content_length': len(content)
                })
            
            return result
            
        except Exception as e:
            return {
                "error": f"Generation failed with {provider_name}: {str(e)}",
                "details": traceback.format_exc()
            }
    
    def create_enhanced_prompt(self, content: str, settings: Dict, topics_str: str, complexity: Dict) -> str:
        """Create enhanced prompt based on content analysis"""
        
        difficulty_guidance = {
            'easy': "Focus on basic recall and simple understanding. Use straightforward language.",
            'medium': "Include analysis and application questions. Moderate complexity.",
            'hard': "Incorporate evaluation, synthesis, and complex reasoning. Advanced vocabulary."
        }
        
        prompt = f"""You are an expert educational assessment creator. Generate diverse, high-quality questions from the provided content.

CONTENT TO ANALYZE:
{content[:2500]}

CONTENT ANALYSIS:
- Reading Level: {complexity['estimated_reading_level']}
- Average Sentence Length: {complexity['avg_sentence_length']} words
- Vocabulary Diversity: {complexity['vocabulary_diversity']}
- Key Topics: {topics_str}

GENERATION REQUIREMENTS:
- Generate {settings['num_questions']} questions total
- Difficulty levels: {', '.join(settings['difficulty_levels'])}
- Question types: {', '.join(settings['question_types'])}
- Include Bloom's taxonomy levels: Remember, Understand, Apply, Analyze, Evaluate, Create

DIFFICULTY GUIDANCE:
{chr(10).join(f"- {level.title()}: {guidance}" for level, guidance in difficulty_guidance.items() if level in settings['difficulty_levels'])}

ENHANCED FEATURES:
- Add realistic distractors for multiple choice questions
- Include prerequisite knowledge indicators
- Provide detailed explanations with examples
- Add follow-up question suggestions
- Include common misconceptions to address

OUTPUT FORMAT (JSON):
{{
  "questions": [
    {{
      "id": "unique_id",
      "difficulty": "easy|medium|hard",
      "difficulty_score": 5,
      "type": "multiple_choice|true_false|short_answer|essay|fill_blank",
      "cognitive_level": "remember|understand|apply|analyze|evaluate|create",
      "question": "Question text",
      "options": ["Option A", "Option B", "Option C", "Option D"],
      "correct_answer": "The correct answer",
      "explanation": "Detailed explanation with reasoning",
      "keywords": ["keyword1", "keyword2"],
      "estimated_time": 3,
      "points": 10,
      "prerequisites": ["required knowledge"],
      "common_mistakes": ["typical errors students make"],
      "follow_up": "Suggested follow-up question or topic"
    }}
  ],
  "metadata": {{
    "total_questions": {settings['num_questions']},
    "content_topics": ["{topics_str}"],
    "difficulty_distribution": {{"easy": 0, "medium": 0, "hard": 0}},
    "cognitive_distribution": {{"remember": 0, "understand": 0, "apply": 0, "analyze": 0, "evaluate": 0, "create": 0}}
  }}
}}"""
        
        return prompt
    
    def ensemble_generation(self, content: str, settings: Dict) -> Dict:
        """Generate questions using multiple providers and merge results"""
        available_models = self.model_provider.get_available_models()
        
        if len(available_models) < 2:
            # Fallback to single provider
            provider_name = available_models[0] if available_models else None
            if not provider_name:
                return {"error": "No providers available"}
            return self.generate_questions_with_provider(content, settings, provider_name)
        
        results = []
        costs = []
        
        selected_providers = settings.get('selected_providers', available_models[:2])
        
        for provider_name in selected_providers:
            if provider_name in available_models:
                result = self.generate_questions_with_provider(content, settings, provider_name)
                if 'error' not in result:
                    results.append(result)
                    costs.append(result.get('usage', {}).get('cost', 0))
        
        if not results:
            return {"error": "All providers failed to generate questions"}
        
        # Merge and deduplicate questions
        return self.merge_ensemble_results(results, costs)
    
    def merge_ensemble_results(self, results: List[Dict], costs: List[float]) -> Dict:
        """Intelligently merge results from multiple providers"""
        all_questions = []
        seen_questions = set()
        
        for result in results:
            for question in result.get('questions', []):
                # Simple deduplication based on question text similarity
                question_key = question['question'].lower().replace(' ', '')[:50]
                if question_key not in seen_questions:
                    seen_questions.add(question_key)
                    all_questions.append(question)
        
        # Sort by quality indicators (difficulty score, explanation length)
        all_questions.sort(key=lambda q: (
            q.get('difficulty_score', 5),
            len(q.get('explanation', '')),
            len(q.get('keywords', []))
        ), reverse=True)
        
        return {
                "questions": all_questions,
                "metadata": {
                    "total_questions": len(all_questions),
                    "providers_used": [r.get('provider', 'unknown') for r in results],
                    "total_cost": sum(costs),
                    "ensemble_used": True,
                    "generation_timestamp": datetime.now().isoformat()
                }
            }
    
    def export_questions(self, questions: List[Dict], format_type: str) -> str:
        """Export questions in various formats"""
        if format_type == 'json':
            return json.dumps(questions, indent=2)
        
        elif format_type == 'csv':
            df = pd.DataFrame(questions)
            return df.to_csv(index=False)
        
        elif format_type == 'yaml':
            return yaml.dump(questions, default_flow_style=False)
        
        elif format_type == 'markdown':
            md_content = "# Generated Questions\n\n"
            for i, q in enumerate(questions, 1):
                md_content += f"## Question {i}\n\n"
                md_content += f"**Type:** {q.get('type', 'N/A')}\n"
                md_content += f"**Difficulty:** {q.get('difficulty', 'N/A')}\n"
                md_content += f"**Points:** {q.get('points', 'N/A')}\n\n"
                md_content += f"**Question:** {q.get('question', 'N/A')}\n\n"
                
                if q.get('options'):
                    md_content += "**Options:**\n"
                    for opt in q['options']:
                        md_content += f"- {opt}\n"
                    md_content += "\n"
                
                md_content += f"**Answer:** {q.get('correct_answer', 'N/A')}\n\n"
                md_content += f"**Explanation:** {q.get('explanation', 'N/A')}\n\n"
                md_content += "---\n\n"
            
            return md_content
        
        return ""
    
    def get_analytics_dashboard(self) -> Dict:
        """Generate analytics dashboard data"""
        if not self.analytics_data:
            return {}
        
        df = pd.DataFrame(self.analytics_data)
        
        analytics = {
            'total_generations': len(df),
            'total_questions': df['questions_generated'].sum(),
            'total_cost': df['cost'].sum(),
            'avg_cost_per_question': df['cost'].sum() / df['questions_generated'].sum() if df['questions_generated'].sum() > 0 else 0,
            'provider_usage': df['provider'].value_counts().to_dict(),
            'cost_by_provider': df.groupby('provider')['cost'].sum().to_dict(),
            'questions_by_provider': df.groupby('provider')['questions_generated'].sum().to_dict()
        }
        
        return analytics

# Streamlit UI
def main():
    st.set_page_config(
        page_title="AI Question Generator",
        page_icon="🤖",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    st.title("🤖 AI Question Generator with LangChain")
    st.markdown("Generate high-quality questions from documents using multiple AI providers")
    
    # Initialize session state
    if 'generator' not in st.session_state:
        st.session_state.generator = AdvancedQuestionGenerator()
    
    if 'questions' not in st.session_state:
        st.session_state.questions = []
    
    if 'generation_history' not in st.session_state:
        st.session_state.generation_history = []
    
    # Sidebar for model configuration
    with st.sidebar:
        st.header("🔧 Model Configuration")
        
        # Provider selection
        provider_tabs = st.tabs(["OpenAI", "Claude", "DeepSeek", "Custom"])
        
        provider_configs = {}
        
        with provider_tabs[0]:
            st.subheader("OpenAI Configuration")
            openai_api_key = st.text_input("OpenAI API Key", type="password", key="openai_key")
            openai_model = st.selectbox("Model", ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo", "gpt-3.5-turbo"], key="openai_model")
            openai_temp = st.slider("Temperature", 0.0, 1.0, 0.7, key="openai_temp")
            
            if openai_api_key:
                provider_configs['openai'] = {
                    'api_key': openai_api_key,
                    'model': openai_model,
                    'temperature': openai_temp,
                    'max_tokens': 3000,
                    'chat_model': True
                }
        
        with provider_tabs[1]:
            st.subheader("Claude Configuration")
            claude_api_key = st.text_input("Claude API Key", type="password", key="claude_key")
            claude_model = st.selectbox("Model", ["claude-3-opus-20240229", "claude-3-sonnet-20240229", "claude-3-haiku-20240307"], key="claude_model")
            claude_temp = st.slider("Temperature", 0.0, 1.0, 0.7, key="claude_temp")
            
            if claude_api_key:
                provider_configs['claude'] = {
                    'api_key': claude_api_key,
                    'model': claude_model,
                    'temperature': claude_temp,
                    'max_tokens': 3000
                }
        
        with provider_tabs[2]:
            st.subheader("DeepSeek Configuration")
            deepseek_api_key = st.text_input("DeepSeek API Key", type="password", key="deepseek_key")
            deepseek_model = st.selectbox("Model", ["deepseek-chat", "deepseek-coder"], key="deepseek_model")
            deepseek_temp = st.slider("Temperature", 0.0, 1.0, 0.7, key="deepseek_temp")
            
            if deepseek_api_key:
                provider_configs['deepseek'] = {
                    'api_key': deepseek_api_key,
                    'model': deepseek_model,
                    'temperature': deepseek_temp,
                    'max_tokens': 3000
                }
        
        with provider_tabs[3]:
            st.subheader("Custom API Configuration")
            custom_api_key = st.text_input("API Key", type="password", key="custom_key")
            custom_base_url = st.text_input("Base URL", placeholder="https://api.example.com/v1", key="custom_url")
            custom_model = st.text_input("Model Name", placeholder="custom-model", key="custom_model")
            custom_provider_name = st.text_input("Provider Name", placeholder="Custom Provider", key="custom_provider")
            custom_cost = st.number_input("Cost per 1K tokens", value=0.01, key="custom_cost")
            custom_temp = st.slider("Temperature", 0.0, 1.0, 0.7, key="custom_temp")
            
            if custom_api_key and custom_base_url and custom_model:
                provider_configs['custom'] = {
                    'api_key': custom_api_key,
                    'base_url': custom_base_url,
                    'model': custom_model,
                    'provider_name': custom_provider_name,
                    'cost_per_1k': custom_cost,
                    'temperature': custom_temp,
                    'max_tokens': 3000
                }
        
        # Initialize providers
        if st.button("Initialize Providers", type="primary"):
            if provider_configs:
                with st.spinner("Initializing AI providers..."):
                    errors = st.session_state.generator.initialize_providers(provider_configs)
                    
                    if errors:
                        for provider, error in errors.items():
                            st.error(f"Error initializing {provider}: {error}")
                    else:
                        st.success("All providers initialized successfully!")
                        st.rerun()
            else:
                st.warning("Please configure at least one provider")
        
        # Show available providers
        available_providers = st.session_state.generator.model_provider.get_available_models()
        if available_providers:
            st.success(f"Active providers: {', '.join(available_providers)}")
        else:
            st.warning("No providers configured")
    
    # Main content area
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("📄 Document Upload")
        
        # File upload
        uploaded_files = st.file_uploader(
            "Upload documents (PDF, DOCX, TXT, CSV, XLSX)",
            type=['pdf', 'docx', 'txt', 'csv', 'xlsx'],
            accept_multiple_files=True
        )
        
        # Text input as alternative
        manual_text = st.text_area("Or paste text directly:", height=200)
        
        if uploaded_files or manual_text:
            st.header("⚙️ Generation Settings")
            
            # Generation settings
            col_settings1, col_settings2 = st.columns(2)
            
            with col_settings1:
                num_questions = st.number_input("Number of Questions", min_value=1, max_value=50, value=10)
                
                difficulty_levels = st.multiselect(
                    "Difficulty Levels",
                    ["easy", "medium", "hard"],
                    default=["easy", "medium"]
                )
                
                question_types = st.multiselect(
                    "Question Types",
                    ["multiple_choice", "true_false", "short_answer", "essay", "fill_blank"],
                    default=["multiple_choice", "short_answer"]
                )
            
            with col_settings2:
                generation_mode = st.selectbox(
                    "Generation Mode",
                    ["Single Provider", "Ensemble (Multiple Providers)"]
                )
                
                if generation_mode == "Single Provider":
                    selected_provider = st.selectbox(
                        "Select Provider",
                        available_providers if available_providers else ["No providers available"]
                    )
                else:
                    selected_providers = st.multiselect(
                        "Select Providers",
                        available_providers if available_providers else ["No providers available"],
                        default=available_providers[:2] if len(available_providers) >= 2 else available_providers
                    )
            
            # Content preview
            if st.button("Preview Content"):
                content = ""
                
                if uploaded_files:
                    for uploaded_file in uploaded_files:
                        file_content = st.session_state.generator.extract_content(uploaded_file)
                        content += f"\n\n--- {uploaded_file.name} ---\n{file_content}"
                
                if manual_text:
                    content += f"\n\n--- Manual Input ---\n{manual_text}"
                
                if content:
                    st.subheader("Content Preview")
                    st.text_area("Extracted Content", content[:2000] + "..." if len(content) > 2000 else content, height=300)
                    
                    # Content analysis
                    analysis = st.session_state.generator.analyze_content_complexity(content)
                    st.subheader("Content Analysis")
                    
                    col_analysis1, col_analysis2, col_analysis3 = st.columns(3)
                    with col_analysis1:
                        st.metric("Total Words", analysis['total_words'])
                        st.metric("Reading Level", analysis['estimated_reading_level'])
                    with col_analysis2:
                        st.metric("Sentences", analysis['total_sentences'])
                        st.metric("Avg Sentence Length", f"{analysis['avg_sentence_length']:.1f}")
                    with col_analysis3:
                        st.metric("Vocabulary Diversity", f"{analysis['vocabulary_diversity']:.3f}")
                        st.metric("Complex Words %", f"{analysis['complex_words_ratio']*100:.1f}%")
            
            # Generate questions
            if st.button("🚀 Generate Questions", type="primary"):
                if not available_providers:
                    st.error("Please configure and initialize at least one AI provider")
                    return
                
                # Prepare content
                content = ""
                
                if uploaded_files:
                    for uploaded_file in uploaded_files:
                        file_content = st.session_state.generator.extract_content(uploaded_file)
                        content += f"\n\n{file_content}"
                
                if manual_text:
                    content += f"\n\n{manual_text}"
                
                if not content.strip():
                    st.error("Please provide content to generate questions from")
                    return
                
                # Prepare settings
                settings = {
                    'num_questions': num_questions,
                    'difficulty_levels': difficulty_levels,
                    'question_types': question_types
                }
                
                if generation_mode == "Ensemble (Multiple Providers)":
                    settings['selected_providers'] = selected_providers
                
                # Generate questions
                with st.spinner("Generating questions... This may take a moment."):
                    try:
                        if generation_mode == "Single Provider":
                            result = st.session_state.generator.generate_questions_with_provider(
                                content, settings, selected_provider
                            )
                        else:
                            result = st.session_state.generator.ensemble_generation(content, settings)
                        
                        if 'error' in result:
                            st.error(f"Generation failed: {result['error']}")
                            if 'details' in result:
                                st.error(f"Details: {result['details']}")
                        else:
                            st.session_state.questions = result.get('questions', [])
                            st.session_state.generation_history.append({
                                'timestamp': datetime.now(),
                                'provider': result.get('provider', 'Unknown'),
                                'questions_count': len(st.session_state.questions),
                                'cost': result.get('usage', {}).get('cost', 0)
                            })
                            
                            st.success(f"Successfully generated {len(st.session_state.questions)} questions!")
                            st.rerun()
                    
                    except Exception as e:
                        st.error(f"An error occurred: {str(e)}")
                        st.error(traceback.format_exc())
    
    with col2:
        st.header("📊 Analytics")
        
        # Usage analytics
        analytics = st.session_state.generator.get_analytics_dashboard()
        
        if analytics:
            st.metric("Total Generations", analytics['total_generations'])
            st.metric("Total Questions", analytics['total_questions'])
            st.metric("Total Cost", f"${analytics['total_cost']:.4f}")
            st.metric("Avg Cost/Question", f"${analytics['avg_cost_per_question']:.4f}")
            
            # Provider usage chart
            if analytics['provider_usage']:
                fig = px.pie(
                    values=list(analytics['provider_usage'].values()),
                    names=list(analytics['provider_usage'].keys()),
                    title="Provider Usage"
                )
                st.plotly_chart(fig, use_container_width=True)
        
        # Cost tracking
        cost_tracking = st.session_state.generator.model_provider.cost_tracking
        if cost_tracking:
            st.subheader("Cost Tracking")
            for provider, data in cost_tracking.items():
                st.write(f"**{provider}:**")
                st.write(f"- Total Cost: ${data['total_cost']:.4f}")
                st.write(f"- Total Tokens: {data['total_tokens']:,}")
    
    # Display generated questions
    if st.session_state.questions:
        st.header("📝 Generated Questions")
        
        # Export options
        col_export1, col_export2 = st.columns(2)
        
        with col_export1:
            export_format = st.selectbox("Export Format", ["json", "csv", "yaml", "markdown"])
        
        with col_export2:
            if st.button("📥 Export Questions"):
                exported_content = st.session_state.generator.export_questions(
                    st.session_state.questions, export_format
                )
                
                st.download_button(
                    label=f"Download {export_format.upper()}",
                    data=exported_content,
                    file_name=f"questions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{export_format}",
                    mime=f"text/{export_format}"
                )
        
        # Question display
        for i, question in enumerate(st.session_state.questions):
            with st.expander(f"Question {i+1} - {question.get('type', 'Unknown').title()} ({question.get('difficulty', 'Unknown')})"):
                st.write(f"**Question:** {question.get('question', 'N/A')}")
                
                if question.get('options'):
                    st.write("**Options:**")
                    for j, option in enumerate(question['options']):
                        st.write(f"{chr(65+j)}. {option}")
                
                st.write(f"**Answer:** {question.get('correct_answer', 'N/A')}")
                st.write(f"**Explanation:** {question.get('explanation', 'N/A')}")
                
                # Additional metadata
                col_meta1, col_meta2, col_meta3 = st.columns(3)
                with col_meta1:
                    st.write(f"**Difficulty Score:** {question.get('difficulty_score', 'N/A')}")
                    st.write(f"**Points:** {question.get('points', 'N/A')}")
                with col_meta2:
                    st.write(f"**Cognitive Level:** {question.get('cognitive_level', 'N/A')}")
                    st.write(f"**Est. Time:** {question.get('estimated_time', 'N/A')} min")
                with col_meta3:
                    if question.get('keywords'):
                        st.write(f"**Keywords:** {', '.join(question['keywords'])}")
                    if question.get('prerequisites'):
                        st.write(f"**Prerequisites:** {', '.join(question['prerequisites'])}")

if __name__ == "__main__":
    main()